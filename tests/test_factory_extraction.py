"""Rule-based (non-LLM) requirement extraction. Pure functions, offline."""

import io
from pathlib import Path

import pytest

from s7_delivery.factory import extraction

EPIC_S7_001 = (Path(__file__).resolve().parent.parent / "requirements" / "epics" / "EPIC-S7-001.md").read_text()

PLAIN_TEXT = """Claims Deductible Handling

Apply policy deductible during claim intake to ensure valid claim processing
and accurate payable amount calculation.

Add a per-policy deductible and apply it during claim intake.

- Policy record must contain a deductible amount.
- Reject claim if claim amount is at or below the policy deductible.
- For accepted claims, calculate payable amount = claim amount - deductible
  and store it.
"""

NO_STRUCTURE_TEXT = "Sponsors need a way to submit claims online without calling support."


def test_extract_from_epic_s7_001_finds_real_title_and_objective():
    result = extraction.extract_requirement(EPIC_S7_001)
    assert result["epic_title"] == "Online disability claim submission for plan sponsors"
    assert "guided online way" in result["business_objective"]
    assert result["requirement_summary"]
    assert len(result["extracted_requirements"]) >= 1
    assert all(r["rule_id"].startswith("REQ-") for r in result["extracted_requirements"])


def test_extract_caps_at_twelve_and_dedupes():
    result = extraction.extract_requirement(EPIC_S7_001)
    assert len(result["extracted_requirements"]) <= 12
    texts = [r["text"] for r in result["extracted_requirements"]]
    assert len(texts) == len(set(texts))


def test_extract_from_plain_text_with_bullets():
    result = extraction.extract_requirement(PLAIN_TEXT)
    assert result["epic_title"] == "Claims Deductible Handling"
    assert "deductible" in result["business_objective"].lower()
    assert len(result["extracted_requirements"]) == 3
    assert result["extracted_requirements"][0]["rule_id"] == "REQ-01"
    assert "deductible amount" in result["extracted_requirements"][0]["text"]


def test_extract_wraps_multiline_bullets_into_one_item():
    result = extraction.extract_requirement(PLAIN_TEXT)
    payable = [r for r in result["extracted_requirements"] if "payable amount" in r["text"]][0]
    assert "and store it" in payable["text"]


def test_extract_falls_back_to_first_line_title_and_trigger_sentences():
    result = extraction.extract_requirement(NO_STRUCTURE_TEXT)
    assert result["epic_title"] == NO_STRUCTURE_TEXT
    assert result["extracted_requirements"]


def test_extract_never_returns_silently_empty_requirements():
    result = extraction.extract_requirement("A short requirement with no lists or trigger words at all here.")
    assert len(result["extracted_requirements"]) == 1
    assert "No discrete requirements detected" in result["extracted_requirements"][0]["text"]


def test_extract_rejects_empty_text():
    with pytest.raises(extraction.ExtractionError, match="empty"):
        extraction.extract_requirement("   ")


def test_decode_txt_and_md():
    assert extraction.decode_source("req.txt", b"Hello world") == "Hello world"
    assert extraction.decode_source("req.md", "café".encode()) == "café"


def test_decode_unsupported_extension_raises():
    with pytest.raises(extraction.ExtractionError, match="Unsupported file type"):
        extraction.decode_source("req.xlsx", b"data")


def test_decode_pdf(monkeypatch):
    class FakePage:
        def __init__(self, text):
            self._text = text

        def extract_text(self):
            return self._text

    class FakeReader:
        def __init__(self, stream):
            self.pages = [FakePage("Page one text."), FakePage("Page two text.")]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)
    text = extraction.decode_source("req.pdf", b"%PDF-fake-bytes")
    assert "Page one text." in text
    assert "Page two text." in text


def test_decode_docx_real_roundtrip():
    import docx

    doc = docx.Document()
    doc.add_heading("Claims Deductible Handling", level=1)
    doc.add_paragraph("Apply the deductible during claim intake.")
    buf = io.BytesIO()
    doc.save(buf)

    text = extraction.decode_source("req.docx", buf.getvalue())
    assert "# Claims Deductible Handling" in text
    assert "Apply the deductible during claim intake." in text
    result = extraction.extract_requirement(text)
    assert result["epic_title"] == "Claims Deductible Handling"


def test_decode_pdf_wraps_parse_failure_as_extraction_error(monkeypatch):
    class BoomReader:
        def __init__(self, stream):
            raise ValueError("not a valid PDF stream")

    monkeypatch.setattr("pypdf.PdfReader", BoomReader)
    with pytest.raises(extraction.ExtractionError, match="req.pdf"):
        extraction.decode_source("req.pdf", b"not really a pdf")


def test_decode_docx_wraps_bad_zip_as_extraction_error():
    with pytest.raises(extraction.ExtractionError, match="req.docx"):
        extraction.decode_source("req.docx", b"not a zip file")


def test_decode_docx_bullet_list_recognized_as_requirement_items():
    import docx

    doc = docx.Document()
    doc.add_heading("Claims Deductible Handling", level=1)
    doc.add_paragraph("Apply the deductible during claim intake.")
    doc.add_paragraph("Policy record contains a deductible amount.", style="List Bullet")
    doc.add_paragraph(
        "Reject a claim at or below the policy deductible.", style="List Bullet"
    )
    buf = io.BytesIO()
    doc.save(buf)

    text = extraction.decode_source("req.docx", buf.getvalue())
    assert "- Policy record contains a deductible amount." in text
    assert "- Reject a claim at or below the policy deductible." in text

    result = extraction.extract_requirement(text)
    texts = [r["text"] for r in result["extracted_requirements"]]
    assert any("Policy record contains a deductible amount" in t for t in texts)
    assert any("Reject a claim at or below the policy deductible" in t for t in texts)


def test_title_scans_past_leading_preamble_for_first_heading():
    text = (
        "Confidential — Draft v2\n\n"
        "# Online claim submission\n\n"
        "Some body text.\n"
    )
    result = extraction.extract_requirement(text)
    assert result["epic_title"] == "Online claim submission"


def test_pdf_lines_reconstruct_into_parseable_markdown():
    """pypdf emits one line per visual line — no blank lines, no markdown.
    pdf_lines_to_markdown rebuilds the structure the block parser expects
    (2026-08-10: the MapleSure requirement PDF parsed as one giant block)."""
    raw = (
        "MS\n"
        "MAPLESURE INSURANCE\n"
        "Group Benefits Operations\n"
        "BUSINESS REQUIREMENT DOCUMENT · PROJECT SCOPE (MULTI-SPRINT)\n"
        "Online disability claim submission for plan sponsors\n"
        "Request ID REQ-2026-114\n"
        "Priority High\n"
        "1. Business context\n"
        "MapleSure sells group disability coverage to plan sponsors — employer organizations that\n"
        "sponsor coverage for their employees.\n"
        "2. Business objective\n"
        "Give plan sponsors a guided online way to submit a disability claim for a member\n"
        "through SponsorConnect, and let them see that it arrived.\n"
        "3. Requirements\n"
        "1. Identify the plan and member. The sponsor identifies whose claim this is from the\n"
        "policy number and member id they already hold.\n"
        "2. Pre-populate what MapleSure already knows. Member and plan details are shown\n"
        "rather than re-keyed.\n"
        "3. Confirm receipt. The sponsor receives a submission reference.\n"
        "4. Out of scope\n"
        "Adjudication of the claim itself.\n"
    )
    md = extraction.pdf_lines_to_markdown(raw)
    result = extraction.extract_requirement(md)
    assert result["epic_title"] == "Online disability claim submission for plan sponsors"
    assert result["business_objective"].startswith("Give plan sponsors a guided online way")
    texts = [r["text"] for r in result["extracted_requirements"]]
    assert len(texts) == 3
    assert texts[0].startswith("Identify the plan and member.")
    assert "policy number and member id they already hold" in texts[0]
    assert texts[2].startswith("Confirm receipt.")
